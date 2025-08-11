import os
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
import tempfile

class DocumentGenerator:
    """Generate .docx documents based on form data"""
    
    def __init__(self):
        self.template_dir = os.path.join(settings.BASE_DIR)
    
    def generate_fax_document(self, form_data, device_type):
        """Generate a fax document based on form data and device type"""
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Fax Document - {device_type.replace("_", " ").title()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add patient information section
        doc.add_heading('Patient Information', level=1)
        
        patient_info = [
            ('Name', form_data.get('name', 'N/A')),
            ('Phone', form_data.get('phone', 'N/A')),
            ('Address', form_data.get('address', 'N/A')),
            ('City', form_data.get('city', 'N/A')),
            ('State', form_data.get('state', 'N/A')),
            ('ZIP', form_data.get('zip', 'N/A')),
            ('Date of Birth', form_data.get('dob', 'N/A')),
            ('Medicare', form_data.get('medicare', 'N/A')),
        ]
        
        for label, value in patient_info:
            if value and value != 'N/A':
                p = doc.add_paragraph()
                p.add_run(f'{label}: ').bold = True
                p.add_run(str(value))
        
        # Add PCP information section
        doc.add_heading('Primary Care Physician (PCP) Information', level=1)
        
        pcp_info = [
            ('Name', form_data.get('pcp_name', 'N/A')),
            ('Address', form_data.get('pcp_address', 'N/A')),
            ('City', form_data.get('pcp_city', 'N/A')),
            ('State', form_data.get('pcp_state', 'N/A')),
            ('ZIP', form_data.get('pcp_zip', 'N/A')),
            ('Phone', form_data.get('pcp_phone', 'N/A')),
            ('Fax', form_data.get('pcp_fax', 'N/A')),
            ('NPI', form_data.get('pcp_npi', 'N/A')),
        ]
        
        for label, value in pcp_info:
            if value and value != 'N/A':
                p = doc.add_paragraph()
                p.add_run(f'{label}: ').bold = True
                p.add_run(str(value))
        
        # Add device information
        doc.add_heading('Device Information', level=1)
        device_paragraph = doc.add_paragraph()
        device_paragraph.add_run('Device Type: ').bold = True
        device_paragraph.add_run(device_type.replace('_', ' ').title())
        
        # Add notes section
        doc.add_heading('Notes', level=1)
        doc.add_paragraph('This document was generated automatically by the CGM Fax Management System.')
        
        # Generate filename
        patient_name = form_data.get('name', 'Unknown').replace(' ', '_')
        device_code = device_type.upper()
        filename = f"{patient_name}-{device_code}.docx"
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)
        doc.save(temp_path)
        
        return temp_path, filename
    
    def generate_from_template(self, form_data, device_type):
        """Generate document from existing template if available"""
        
        # Map device types to template files
        template_map = {
            'cgm': 'do.docx',
            'ankle': 'docs_braces/Ankle_DO.docx',
            'knee': 'docs_braces/Knee_DO.docx',
            'back': 'docs_braces/Back_DO.docx',
            'hip': 'docs_braces/Hip_DO.docx',
            'shoulder': 'docs_braces/Shoulder_DO.docx',
            'wrist': 'docs_braces/Wrist_DO.docx',
        }
        
        template_file = template_map.get(device_type)
        if not template_file:
            # Fall back to generating new document
            return self.generate_fax_document(form_data, device_type)
        
        template_path = os.path.join(self.template_dir, template_file)
        
        if not os.path.exists(template_path):
            # Template doesn't exist, generate new document
            return self.generate_fax_document(form_data, device_type)
        
        try:
            # Load existing template
            doc = Document(template_path)
            
            # Replace placeholders in the template
            self._replace_placeholders(doc, form_data, device_type)
            
            # Generate filename
            patient_name = form_data.get('name', 'Unknown').replace(' ', '_')
            device_code = device_type.upper()
            filename = f"{patient_name}-{device_code}.docx"
            
            # Save to temporary file
            temp_dir = tempfile.gettempdir()
            temp_path = os.path.join(temp_dir, filename)
            doc.save(temp_path)
            
            return temp_path, filename
            
        except Exception as e:
            # If template processing fails, fall back to generating new document
            print(f"Template processing failed: {e}")
            return self.generate_fax_document(form_data, device_type)
    
    def _replace_placeholders(self, doc, form_data, device_type):
        """Replace placeholders in the template document"""
        
        # Define replacement mappings
        replacements = {
            '{{name}}': form_data.get('name', 'N/A'),
            '{{phone}}': form_data.get('phone', 'N/A'),
            '{{address}}': form_data.get('address', 'N/A'),
            '{{city}}': form_data.get('city', 'N/A'),
            '{{state}}': form_data.get('state', 'N/A'),
            '{{zip}}': form_data.get('zip', 'N/A'),
            '{{dob}}': str(form_data.get('dob', 'N/A')),
            '{{medicare}}': form_data.get('medicare', 'N/A'),
            '{{insurance}}': form_data.get('medicare', 'N/A'),  # Map medicare to insurance
            '{{pcp_name}}': form_data.get('pcp_name', 'N/A'),
            '{{pcp_address}}': form_data.get('pcp_address', 'N/A'),
            '{{pcp_city}}': form_data.get('pcp_city', 'N/A'),
            '{{pcp_state}}': form_data.get('pcp_state', 'N/A'),
            '{{pcp_zip}}': form_data.get('pcp_zip', 'N/A'),
            '{{pcp_phone}}': form_data.get('pcp_phone', 'N/A'),
            '{{pcp_fax}}': form_data.get('pcp_fax', 'N/A'),
            '{{pcp_npi}}': form_data.get('pcp_npi', 'N/A'),
            '{{date}}': str(datetime.date.today()),  # Add current date
            '{{cgm}}': device_type.replace('_', ' ').title(),
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
